import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
IMAGE_PATTERN = re.compile(r"^!\[(.*?)\]\(([^)]+)\)(?:\{[^}]*\})?$")
HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.*)$")
LIST_PATTERN = re.compile(r"^(\d+\.|[-*])\s+(.*)$")


def parse_args():
    parser = argparse.ArgumentParser(description="Build the Chinese dissertation Word draft from Markdown")
    parser.add_argument("--input", default="博士学位论文初稿.md")
    parser.add_argument("--output", default="博士学位论文初稿.docx")
    parser.add_argument("--title", default="乳腺超声图像渐进分割方法研究")
    return parser.parse_args()


def set_east_asia_font(run, name):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(clean_inline(text))
    run.bold = bold
    run.font.size = Pt(9)
    set_east_asia_font(run, "宋体")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def add_toc(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "请在 Word 中右键更新目录"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, placeholder, end))


def clean_inline(text):
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1 (\2)", text)
    return text.strip()


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.3)
    section.footer_distance = Cm(1.3)
    add_page_number(section.footer.paragraphs[0])

    normal = document.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.space_after = Pt(0)
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")

    heading_sizes = {1: 18, 2: 15, 3: 13, 4: 11}
    for level, size in heading_sizes.items():
        style = document.styles["Heading {}".format(level)]
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.first_line_indent = Cm(0)


def add_markdown_paragraph(document, text, style=None, quote=False):
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.first_line_indent = Cm(0) if style or quote else Cm(0.74)
    if quote:
        paragraph.paragraph_format.left_indent = Cm(0.8)
        paragraph.paragraph_format.right_indent = Cm(0.8)
    run = paragraph.add_run(clean_inline(text))
    set_east_asia_font(run, "宋体")
    if quote:
        run.font.color.rgb = RGBColor(80, 80, 80)
    return paragraph


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    return rows, index


def add_table(document, rows):
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, values in enumerate(rows):
        for column_index in range(columns):
            value = values[column_index] if column_index < len(values) else ""
            set_cell_text(table.cell(row_index, column_index), value, bold=row_index == 0)
    if rows:
        set_repeat_table_header(table.rows[0])
    document.add_paragraph()


def add_image(document, caption, image_path):
    path = (ROOT / image_path).resolve()
    if not path.is_file():
        add_markdown_paragraph(document, "[缺失图片：{}]".format(image_path))
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(15.0))
    if caption:
        caption_paragraph = document.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_paragraph.add_run(caption)
        caption_run.font.size = Pt(9)
        set_east_asia_font(caption_run, "宋体")


def build_document(markdown_path, output_path, title):
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_document(document)
    document.core_properties.title = title
    document.core_properties.subject = "U-Net, BUSI, 医学图像分割"

    index = 0
    in_code = False
    code_lines = []
    inserted_toc = False
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if stripped.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Cm(0.6)
                paragraph.paragraph_format.first_line_indent = Cm(0)
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not stripped or stripped == "---":
            index += 1
            continue

        heading = HEADING_PATTERN.match(stripped)
        if heading:
            level = min(len(heading.group(1)), 4)
            title = clean_inline(heading.group(2))
            paragraph = document.add_heading(title, level=level)
            paragraph.paragraph_format.first_line_indent = Cm(0)
            if title == "目录" and not inserted_toc:
                add_toc(document.add_paragraph())
                inserted_toc = True
            index += 1
            continue

        image = IMAGE_PATTERN.match(stripped)
        if image:
            add_image(document, image.group(1), image.group(2).split()[0])
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue

        list_item = LIST_PATTERN.match(stripped)
        if list_item:
            style = "List Number" if list_item.group(1).endswith(".") else "List Bullet"
            add_markdown_paragraph(document, list_item.group(2), style=style)
            index += 1
            continue

        if stripped.startswith(">"):
            add_markdown_paragraph(document, stripped.lstrip("> "), quote=True)
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if (not candidate or candidate.startswith(("#", "|", "!", ">", "```"))
                    or LIST_PATTERN.match(candidate)):
                break
            paragraph_lines.append(candidate)
            index += 1
        add_markdown_paragraph(document, " ".join(paragraph_lines))

    document.save(output_path)


def main():
    args = parse_args()
    input_path = (ROOT / args.input).resolve()
    output_path = (ROOT / args.output).resolve()
    build_document(input_path, output_path, args.title)
    print("wrote {} ({} bytes)".format(output_path, output_path.stat().st_size))


if __name__ == "__main__":
    main()
