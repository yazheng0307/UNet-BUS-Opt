import json
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parent
MARKDOWN = ROOT / "硕士学位论文初稿_低专业度版.md"
WORD = ROOT / "硕士学位论文初稿_低专业度版.docx"
SIGNIFICANCE = ROOT / "thesis_artifacts/tables/best_seed_significance.json"


def main():
    markdown = MARKDOWN.read_text(encoding="utf-8")
    document = Document(WORD)
    document_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    with ZipFile(WORD) as archive:
        media = [
            name for name in archive.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        ]
    image_paths = re.findall(r"^!\[[^]]*\]\(([^) ]+)", markdown, re.MULTILINE)
    best_seed_metrics = (
        ("72.445", "80.718"),
        ("74.185", "82.256"),
        ("74.193", "82.537"),
        ("74.847", "82.943"),
        ("74.966", "83.089"),
        ("75.442", "83.382"),
    )
    plain_names = (
        "多特征增强模块",
        "全局关系增强模块",
        "边界细化模块",
        "易错区域修正模块",
    )
    mainstream_models = (
        "U-Net",
        "Attention U-Net",
        "U-Net++",
        "U-Net3+",
        "TransUnet",
        "MedT",
        "SwinUnet",
        "UNeXt",
        "CMU-Net",
        "CMUNeXt",
        "Mobile U-ViT",
    )
    chapter3_comparison = markdown[
        markdown.index("## 3.5 "):markdown.index("## 3.6 ")
    ]
    chapter4_comparison = markdown[
        markdown.index("## 4.5 "):markdown.index("## 4.6 ")
    ]
    significance = json.loads(SIGNIFICANCE.read_text(encoding="utf-8"))
    checks = {
        "markdown_exists": MARKDOWN.is_file(),
        "word_exists": WORD.is_file(),
        "four_plain_names": all(
            name in markdown and name in document_text for name in plain_names
        ),
        "old_names_removed": not any(
            name in markdown for name in ("TriCAR", "CGR-Bridge", "BD-CoRefine", "UDER")
        ),
        "all_best_seed_metrics_present": all(
            iou in markdown and dice in markdown
            and iou in document_text and dice in document_text
            for iou, dice in best_seed_metrics
        ),
        "ten_markdown_images_exist": (
            len(image_paths) == 10
            and all((ROOT / path).is_file() for path in image_paths)
        ),
        "ten_word_images": len(media) == 10,
        "toc_heading": any(
            paragraph.text == "目录" for paragraph in document.paragraphs
        ),
        "no_missing_image_marker": "[缺失图片：" not in document_text,
        "equation_delimiters_balanced": markdown.count("$$") % 2 == 0,
        "title_matches": document.core_properties.title == (
            "基于改进 U-Net 的乳腺超声图像分割方法研究"
        ),
        "mainstream_review_contains_eleven_models": (
            "### 1.3.2 主流对比模型介绍" in markdown
            and all(model in markdown for model in mainstream_models)
        ),
        "both_research_chapters_have_comparisons": (
            "## 3.5 与主流模型的参考对比" in markdown
            and "## 4.5 与主流模型的参考对比" in markdown
            and "UnetAB（本文，seed 41）" in markdown
            and "UABCD（本文，seed 73）" in markdown
        ),
        "mobile_uvit_removed_from_comparison_experiments": (
            "Mobile U-ViT" not in chapter3_comparison
            and "Mobile U-ViT" not in chapter4_comparison
        ),
        "reference_metrics_match_source_table": all(
            value in markdown
            for value in (
                "68.61",
                "76.97",
                "71.56",
                "79.86",
                "650.48",
                "199.74 GFLOPs",
            )
        ),
        "fixed_split3_source_confirmation_disclosed": (
            "经原仓库作者确认" in markdown
            and "“±”前的数值对应固定 split 3 的训练结果" in markdown
        ),
        "best_of_three_selection_is_consistent": (
            "按照逐病例平均 IoU 从三个种子中选择最高的一次" in markdown
            and "不分别挑选两项指标的最大值" in markdown
            and "seed 41 的 Dice 为 83.480%" in markdown
            and "不把不同种子的最大 IoU 和最大 Dice 拼接" in markdown
        ),
        "descriptive_difference_analysis_present": (
            "### 3.5.1 差值描述性分析" in markdown
            and "### 4.5.1 差值描述性分析" in markdown
            and "3.287 个 IoU 百分点" in markdown
            and "3.882 个 IoU 百分点" in markdown
            and "| CMUNeXt | +3.287 | +3.083 |" in markdown
            and "| CMUNeXt | +3.882 | +3.522 |" in markdown
        ),
        "module_improvement_explanations_present": (
            "### 3.5.3 为什么 A 和 B 能够带来提升" in markdown
            and "### 4.5.3 为什么 C 和 D 能够继续提升结果" in markdown
        ),
        "best_seed_significance_artifacts_valid": (
            SIGNIFICANCE.is_file()
            and len(significance) == 8
            and all(record["cases"] == 195 for record in significance)
            and "0.0211" in markdown
            and "0.0038" in markdown
        ),
        "new_references_present": all(
            "[{}]".format(index) in markdown for index in range(31, 37)
        ),
        "at_least_fourteen_tables": len(document.tables) >= 14,
    }
    report = {
        "passed": all(checks.values()),
        "checks": checks,
        "document": {
            "paragraphs": len(document.paragraphs),
            "headings": sum(
                paragraph.style.name.startswith("Heading")
                for paragraph in document.paragraphs
            ),
            "tables": len(document.tables),
            "images": len(media),
            "markdown_chars": len(markdown),
            "word_paragraph_chars": sum(
                len(paragraph.text) for paragraph in document.paragraphs
            ),
        },
    }
    print(json.dumps(report, ensure_ascii=False, indent=2))
    if not report["passed"]:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
